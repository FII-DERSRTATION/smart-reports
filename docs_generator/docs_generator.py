from docx import Document

def generate_dummy_docx_file(firstName, lastName, SID,  statsBasedDiagnostic, befundeLabor, medicine_entrance,
                             medicine_exit, anamnesisGen, berteilung):

    document = Document()

    document.add_heading(firstName + ' ' + lastName + ' ' + SID, 0)

    document.add_paragraph('Pregenerated results')

    document.add_heading('Diagnostic: ' + statsBasedDiagnostic, level=1)

    document.add_heading('Befunde labor', level=1)
    document.add_paragraph(befundeLabor, style='Intense Quote')

    document.add_heading('Anamnesis', level=1)
    document.add_paragraph(str(anamnesisGen.strip()).replace('<pad>', ''), style='Intense Quote')

    document.add_heading('Berteilung', level=1)
    document.add_paragraph(str(berteilung.strip()).replace('<pad>', ''), style='Intense Quote')

    document.add_heading('Medicine entrance recommendations', level=1)
    document.add_paragraph(' \n'.join(medicine_entrance[0:10]), style='Intense Quote')

    document.add_heading('Medicine entrance recommendations', level=1)
    document.add_paragraph(' \n'.join(medicine_exit[0:10]), style='Intense Quote')

    document.add_page_break()

    document.save('report.docx')


if __name__ == "__main__":
    generate_dummy_docx_file()